"""Spec Agent — generates functional specification (.md + .docx) from any input."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rich.console import Console

from orchestrator.base_agent import BaseAgent
from orchestrator.settings import settings
from tools.notion import NotionClient

console = Console(legacy_windows=False)

DOCS_DIR = Path("documents/specs")

_ORIGEM_MAP = {
    "text": "Texto Livre",
    "file": "Documento",
    "data": "Documento",
    "auto": "Documento",
}


class SpecAgent(BaseAgent):
    name = "Spec Agent"
    role = "Analista de Especificacao Funcional"
    model: str = settings.architect_model
    prompt_file = "agents/prompts/spec.md"

    def __init__(self) -> None:
        self.model = settings.architect_model
        super().__init__()
        self.notion = NotionClient()

    def run(
        self,
        input_file: str = "",
        input_data: dict[str, Any] | None = None,
        input_text: str = "",
        pipeline: str = "",
        origem: str = "",
    ) -> dict[str, Any]:
        """Generate functional spec (.md + .docx) from any input source.

        Priority: input_text > input_data > input_file > auto-detect latest output.
        """
        console.rule("[bold]Spec Agent")

        # 1. Load input
        raw_input = ""
        if input_text:
            data = {"titulo": input_text[:80], "_raw_text": input_text}
            label = input_text[:40].replace(" ", "_").lower()
            raw_input = input_text
            _origem = origem or "Texto Livre"
        elif input_data:
            data = input_data
            label = data.get("titulo", data.get("sprint", "spec")).replace(" ", "_").lower()
            _origem = origem or "Documento"
        elif input_file:
            content = Path(input_file).read_text(encoding="utf-8")
            try:
                data = json.loads(content)
            except json.JSONDecodeError:
                data = {"titulo": Path(input_file).stem, "_raw_text": content}
                raw_input = content
            label = Path(input_file).stem
            _origem = origem or "Documento"
        else:
            raw, label = self._load_latest_output()
            data = json.loads(raw)
            _origem = origem or "Documento"

        # 2. Detect pipeline
        if not pipeline:
            pipeline = "CWI" if ("epicos" in data or "origem" in data) else "Expansão AI"

        # 3. Extract items or use raw text directly
        raw_text = data.get("_raw_text", "")
        if raw_text:
            items = [{"titulo": data.get("titulo", "Spec"), "_raw_text": raw_text}]
        else:
            items = self._extract_items(data)

        doc_title = data.get("titulo", data.get("origem", label))
        console.print(f"[dim]Pipeline: {pipeline} | Itens: {len(items)}[/]")

        # 4. Generate spec per item
        all_specs: list[dict[str, Any]] = []
        for i, item in enumerate(items, 1):
            console.print(f"[dim]Gerando spec {i}/{len(items)}: {item.get('titulo', item.get('id', '?'))}[/]")
            spec = self._spec_for_item(item, pipeline, doc_title)
            all_specs.append(spec)

        # 5. Merge
        merged = self._merge_specs(all_specs, doc_title, pipeline)

        # 6. Save files
        today = date.today().strftime("%Y_%m_%d")
        slug = label[:40].replace("/", "_")
        base_name = f"spec_{slug}_{today}"

        md_path = self._save_markdown(merged, base_name)
        docx_path = self._save_docx(merged, base_name)

        console.print(f"[green]Markdown:[/] {md_path}")
        console.print(f"[green]Word:[/]     {docx_path}")

        # 7. Persist to Notion Specs DB
        notion_page_id = ""
        if settings.notion_specs_db_id:
            try:
                notion_page_id = self.notion.create_spec_page(
                    spec=merged,
                    input_text=raw_input or json.dumps(data, ensure_ascii=False)[:1000],
                    pipeline=pipeline,
                    origem=_origem,
                )
                console.print(f"[green]Notion:[/]   {notion_page_id}")
            except Exception as e:
                console.print(f"[yellow]Notion falhou (nao critico): {e}[/]")

        # 8. Generate diagrams automatically
        try:
            from orchestrator.agents.diagram_agent import DiagramAgent  # noqa: PLC0415
            diagram_paths = DiagramAgent().run(spec=merged, label=slug)
        except Exception as e:
            console.print(f"[yellow]DiagramAgent falhou (nao critico): {e}[/]")
            diagram_paths = []

        return {
            "md": md_path,
            "docx": docx_path,
            "spec": merged,
            "notion_page_id": notion_page_id,
            "diagrams": diagram_paths,
        }

    # ── Item extraction ───────────────────────────────────────────────────────

    def _extract_items(self, data: dict[str, Any]) -> list[dict[str, Any]]:
        """Extract epics (CWI) or PRDs (Expansão AI) from agent output."""
        if "epicos" in data:
            return data["epicos"]
        if "prds" in data:
            return data["prds"]
        if "features" in data:
            return data["features"]
        # Fallback: treat entire data as single item
        return [data]

    # ── Per-item spec generation ──────────────────────────────────────────────

    def _spec_for_item(self, item: dict[str, Any], pipeline: str, context: str) -> dict[str, Any]:
        user_message = f"""Gere a especificacao funcional para o item abaixo.
Contexto do projeto: {context}
Pipeline: {pipeline}

REGRAS DE CONCISAO (obrigatorio):
- Maximo 3 casos de uso por item
- Fluxo principal: maximo 6 passos
- Maximo 1 fluxo alternativo por caso de uso, com maximo 3 passos
- Maximo 5 regras de negocio
- Maximo 3 requisitos nao funcionais
- Maximo 5 itens no glossario
- Campos de texto: maximo 2 frases cada

ITEM:
{json.dumps(item, ensure_ascii=False, indent=2)}

Retorne apenas o JSON da especificacao, sem texto adicional."""

        response_text = self._run(user_message, max_tokens=8192)
        return self._parse_json_output(response_text)

    # ── Merge ─────────────────────────────────────────────────────────────────

    def _merge_specs(self, specs: list[dict[str, Any]], title: str, pipeline: str) -> dict[str, Any]:
        if len(specs) == 1:
            return specs[0]

        merged: dict[str, Any] = {
            "titulo": title,
            "versao": "1.0",
            "data": date.today().isoformat(),
            "pipeline": pipeline,
            "objetivo": specs[0].get("objetivo", ""),
            "escopo": {"inclui": [], "exclui": []},
            "atores": [],
            "casos_de_uso": [],
            "regras_de_negocio": [],
            "requisitos_nao_funcionais": [],
            "glossario": [],
            "perguntas_em_aberto": [],
        }

        seen_atores: set[str] = set()
        uc_counter = 1
        rn_counter = 1

        for spec in specs:
            for inc in spec.get("escopo", {}).get("inclui", []):
                merged["escopo"]["inclui"].append(inc)
            for exc in spec.get("escopo", {}).get("exclui", []):
                merged["escopo"]["exclui"].append(exc)
            for ator in spec.get("atores", []):
                if ator["ator"] not in seen_atores:
                    merged["atores"].append(ator)
                    seen_atores.add(ator["ator"])
            for uc in spec.get("casos_de_uso", []):
                uc["id"] = f"UC-{uc_counter:03d}"
                uc_counter += 1
                merged["casos_de_uso"].append(uc)
            for rn in spec.get("regras_de_negocio", []):
                rn["id"] = f"RN-{rn_counter:03d}"
                rn_counter += 1
                merged["regras_de_negocio"].append(rn)
            merged["requisitos_nao_funcionais"].extend(spec.get("requisitos_nao_funcionais", []))
            merged["glossario"].extend(spec.get("glossario", []))
            merged["perguntas_em_aberto"].extend(spec.get("perguntas_em_aberto", []))

        return merged

    # ── Markdown ──────────────────────────────────────────────────────────────

    def _save_markdown(self, spec: dict[str, Any], base_name: str) -> Path:
        lines: list[str] = []

        lines.append(f"# {spec.get('titulo', 'Especificacao Funcional')}\n")
        lines.append(f"**Versão:** {spec.get('versao', '1.0')} | "
                     f"**Data:** {spec.get('data', date.today().isoformat())} | "
                     f"**Pipeline:** {spec.get('pipeline', '-')}\n")

        lines.append(f"## Objetivo\n\n{spec.get('objetivo', '')}\n")

        escopo = spec.get("escopo", {})
        if escopo.get("inclui") or escopo.get("exclui"):
            lines.append("## Escopo\n")
            lines.append("**Inclui:**")
            for item in escopo.get("inclui", []):
                lines.append(f"- {item}")
            lines.append("\n**Exclui:**")
            for item in escopo.get("exclui", []):
                lines.append(f"- {item}")
            lines.append("")

        atores = spec.get("atores", [])
        if atores:
            lines.append("## Atores\n")
            for a in atores:
                lines.append(f"**{a['ator']}** — {a['descricao']}")
            lines.append("")

        ucs = spec.get("casos_de_uso", [])
        if ucs:
            lines.append("## Casos de Uso\n")
            for uc in ucs:
                lines.append(f"### {uc['id']} — {uc['nome']}\n")
                lines.append(f"**Ator principal:** {uc.get('ator_principal', '-')}")
                pre = uc.get("pre_condicoes", [])
                if pre:
                    lines.append(f"**Pré-condições:** {'; '.join(pre)}")
                lines.append("\n**Fluxo principal:**")
                for step in uc.get("fluxo_principal", []):
                    lines.append(f"{step}")
                for alt in uc.get("fluxos_alternativos", []):
                    lines.append(f"\n**Alternativo — {alt['condicao']}:**")
                    for step in alt.get("passos", []):
                        lines.append(f"{step}")
                pos = uc.get("pos_condicoes", [])
                if pos:
                    lines.append(f"\n**Pós-condições:** {'; '.join(pos)}")
                lines.append("")

        rns = spec.get("regras_de_negocio", [])
        if rns:
            lines.append("## Regras de Negócio\n")
            for rn in rns:
                lines.append(f"**{rn['id']}** — {rn['regra']}")
                lines.append(f"*Origem: {rn.get('origem', '-')}*\n")

        rnfs = spec.get("requisitos_nao_funcionais", [])
        if rnfs:
            lines.append("## Requisitos Não Funcionais\n")
            for r in rnfs:
                lines.append(f"- **{r['categoria']}:** {r['requisito']}")
            lines.append("")

        glossario = spec.get("glossario", [])
        if glossario:
            lines.append("## Glossário\n")
            for g in glossario:
                lines.append(f"**{g['termo']}:** {g['definicao']}")
            lines.append("")

        perguntas = spec.get("perguntas_em_aberto", [])
        if perguntas:
            lines.append("## Perguntas em Aberto\n")
            for p in perguntas:
                lines.append(f"- {p}")

        out_dir = Path(settings.output_dir) / "specs"
        out_dir.mkdir(parents=True, exist_ok=True)
        path = out_dir / f"{base_name}.md"
        path.write_text("\n".join(lines), encoding="utf-8")
        return path

    # ── DOCX ─────────────────────────────────────────────────────────────────

    def _save_docx(self, spec: dict[str, Any], base_name: str) -> Path:
        doc = Document()
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(11)

        h = doc.add_heading(spec.get("titulo", "Especificacao Funcional"), level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta = doc.add_paragraph(
            f"Versão: {spec.get('versao', '1.0')}   |   "
            f"Data: {spec.get('data', date.today().isoformat())}   |   "
            f"Pipeline: {spec.get('pipeline', '-')}"
        )
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        doc.add_heading("Objetivo", level=1)
        doc.add_paragraph(spec.get("objetivo", ""))

        escopo = spec.get("escopo", {})
        if escopo.get("inclui") or escopo.get("exclui"):
            doc.add_heading("Escopo", level=1)
            doc.add_heading("Inclui", level=2)
            for item in escopo.get("inclui", []):
                doc.add_paragraph(item, style="List Bullet")
            doc.add_heading("Exclui", level=2)
            for item in escopo.get("exclui", []):
                doc.add_paragraph(item, style="List Bullet")

        atores = spec.get("atores", [])
        if atores:
            doc.add_heading("Atores", level=1)
            for a in atores:
                p = doc.add_paragraph()
                p.add_run(f"{a['ator']}: ").bold = True
                p.add_run(a["descricao"])

        ucs = spec.get("casos_de_uso", [])
        if ucs:
            doc.add_heading("Casos de Uso", level=1)
            for uc in ucs:
                doc.add_heading(f"{uc['id']} — {uc['nome']}", level=2)
                p = doc.add_paragraph()
                p.add_run("Ator principal: ").bold = True
                p.add_run(uc.get("ator_principal", "-"))
                pre = uc.get("pre_condicoes", [])
                if pre:
                    p2 = doc.add_paragraph()
                    p2.add_run("Pré-condições: ").bold = True
                    p2.add_run("; ".join(pre))
                doc.add_heading("Fluxo principal", level=3)
                for step in uc.get("fluxo_principal", []):
                    doc.add_paragraph(step, style="List Number")
                for alt in uc.get("fluxos_alternativos", []):
                    doc.add_heading(f"Alternativo: {alt['condicao']}", level=3)
                    for step in alt.get("passos", []):
                        doc.add_paragraph(step, style="List Number")
                pos = uc.get("pos_condicoes", [])
                if pos:
                    p3 = doc.add_paragraph()
                    p3.add_run("Pós-condições: ").bold = True
                    p3.add_run("; ".join(pos))

        rns = spec.get("regras_de_negocio", [])
        if rns:
            doc.add_heading("Regras de Negócio", level=1)
            for rn in rns:
                p = doc.add_paragraph()
                p.add_run(f"{rn['id']} — ").bold = True
                p.add_run(rn["regra"])
                doc.add_paragraph(f"Origem: {rn.get('origem', '-')}", style="List Bullet")

        rnfs = spec.get("requisitos_nao_funcionais", [])
        if rnfs:
            doc.add_heading("Requisitos Não Funcionais", level=1)
            for r in rnfs:
                p = doc.add_paragraph()
                p.add_run(f"{r['categoria']}: ").bold = True
                p.add_run(r["requisito"])

        glossario = spec.get("glossario", [])
        if glossario:
            doc.add_heading("Glossário", level=1)
            for g in glossario:
                p = doc.add_paragraph()
                p.add_run(f"{g['termo']}: ").bold = True
                p.add_run(g["definicao"])

        perguntas = spec.get("perguntas_em_aberto", [])
        if perguntas:
            doc.add_heading("Perguntas em Aberto", level=1)
            for p_text in perguntas:
                doc.add_paragraph(p_text, style="List Bullet")

        DOCS_DIR.mkdir(parents=True, exist_ok=True)
        path = DOCS_DIR / f"{base_name}.docx"
        doc.save(str(path))
        return path

    # ── Fallback loader ───────────────────────────────────────────────────────

    @staticmethod
    def _load_latest_output() -> tuple[str, str]:
        out = Path("outputs/")
        for prefix in ("pm_prds_", "cwi/product_", "cwi/executive_report_"):
            files = sorted(out.glob(f"{prefix}*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
            if files:
                return files[0].read_text(encoding="utf-8"), files[0].stem
        raise FileNotFoundError("Nenhum output de PM ou Product Agent encontrado.")
